"""把 Markdown 设计文档转换为格式化 .docx。"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "./设备接入识别与阻断控制系统_设计文档.docx"
doc = Document()

# 字体与正文样式
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# 页面边距
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.2)


def set_cn_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_h(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13}
    set_cn_font(run, size=sizes.get(level, 12), bold=True,
                color=RGBColor(0x1F, 0x49, 0x82) if level <= 2 else RGBColor(0x33, 0x33, 0x33))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_cn_font(run)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_cn_font(run)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # 设置浅灰背景
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd)


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, bold=True)
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_cn_font(run)
    return t


# ================== 封面 ==================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("设备接入识别与阻断控制系统")
set_cn_font(run, size=24, bold=True, color=RGBColor(0x1F, 0x49, 0x82))
title.paragraph_format.space_before = Pt(80)
title.paragraph_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("—— 总体框架设计与任务需求文档 ——")
set_cn_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))
subtitle.paragraph_format.space_after = Pt(40)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("技术栈:Python (Scapy + FastAPI) · React (Ant Design) · SQLite\n版本:v0.1   ·   日期:2026-05-16")
set_cn_font(run, size=11)

doc.add_page_break()


# ================== 一、项目概述 ==================
add_h("一、项目概述", 1)

add_h("1.1 项目背景", 2)
add_p('随着物联网与移动设备的普及,家庭和小型办公网络中接入的设备越来越多。未经授权的设备(如蹭网设备、可疑终端)可能带来数据泄漏、内网渗透、带宽占用等安全风险。本项目设计并实现一个轻量级的"设备接入识别与阻断控制系统",通过主动扫描 + 被动监听的方式发现局域网内所有在线设备,建立设备画像,识别异常接入行为,并提供阻断与审计能力。')

add_h("1.2 系统目标", 2)
add_p("系统需要达成以下五个核心目标:第一,自动发现局域网内所有在线设备并维护设备资产清单;第二,基于 IP/MAC/主机名/厂商 OUI 等信息建立设备身份画像;第三,通过白名单、黑名单与规则引擎识别异常接入;第四,对异常或黑名单设备执行逻辑阻断(ARP 欺骗 / iptables / 路由器联动);第五,提供可视化界面和完整审计日志,满足追溯需求。")

add_h("1.3 适用场景", 2)
add_p("家庭无线网络的访客管理与防蹭网;小型办公网络的资产盘点与访问控制;实验室、教学场景的网络安全演示。")


# ================== 二、总体框架设计 ==================
add_h("二、总体框架设计", 1)

add_h("2.1 系统分层架构", 2)
add_p("系统采用经典的四层架构:采集层 → 数据层 → 业务逻辑层 → 表示层。", bold=True)
add_table(
    ["层次", "主要职责", "关键技术"],
    [
        ["表示层 (Presentation)", "可视化展示、用户交互、报表导出", "React 18 + Ant Design 5 + ECharts"],
        ["业务逻辑层 (Service)",   "设备识别、风险判定、策略执行、审计", "FastAPI + APScheduler + Pydantic"],
        ["数据层 (Data)",          "设备资产、日志、规则、配置的持久化", "SQLite + SQLAlchemy ORM"],
        ["采集层 (Collector)",     "网络扫描、流量监听、ARP/DHCP 嗅探", "Scapy + python-nmap + psutil"],
    ],
)

add_h("2.2 核心模块划分", 2)
add_p("系统划分为 8 个核心模块,各自承担明确职责且通过定义良好的接口协作。")

modules = [
    ("(1) 设备发现模块 (Discovery)", "负责主动扫描局域网,采用 ARP 扫描 + ICMP Ping + Nmap 端口探测组合手段。支持周期性扫描(默认 30s)和按需触发扫描。"),
    ("(2) 设备指纹模块 (Fingerprint)", "基于 MAC OUI 数据库识别厂商,通过 DHCP Option 55、TTL、TCP 窗口大小等被动指纹识别操作系统,通过 mDNS/NetBIOS/UPnP 获取主机名。"),
    ("(3) 设备登记模块 (Registry)", "维护设备资产表,支持白名单、黑名单、未知三类分类。新发现设备默认进入未知列表等待人工审核。"),
    ("(4) 数据采集模块 (Telemetry)", "记录每台设备的 IP、MAC、主机名、首次上线、最后在线、累计在线时长、流量摘要、访问目标 IP/域名 摘要。"),
    ("(5) 风险判定模块 (Risk Engine)", "基于可配置的规则集判定风险等级(低/中/高/严重)。规则示例包括:陌生 MAC 出现在非工作时间、设备 MAC 变化频繁、扫描行为(端口扫描特征)、与黑名单 IP 通信等。"),
    ("(6) 阻断控制模块 (Blocker)", "支持三种阻断策略 —— ARP 欺骗(向目标设备发送伪造网关 ARP)、iptables/nftables 规则(适用于网关部署)、路由器 API 联动(适用于支持 OpenWRT 的环境)。"),
    ("(7) 审计日志模块 (Audit)", "记录所有接入事件、告警事件、放行操作、阻断操作,字段包括时间戳、设备标识、操作类型、操作人、原因。"),
    ("(8) 可视化模块 (Dashboard)", "提供设备列表、风险概览、实时告警、审计回放、规则配置等界面。"),
]
for name, desc in modules:
    p = doc.add_paragraph()
    r1 = p.add_run(name + ":")
    set_cn_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_cn_font(r2)

add_h("2.3 数据流设计", 2)
add_p("系统运行时的核心数据流如下:")
flow = [
    "扫描线程周期性调用 Discovery 扫描局域网,得到 (IP, MAC) 对列表;",
    "列表交由 Fingerprint 补全厂商、OS、主机名等画像信息;",
    "Registry 比对资产表,标记设备状态为「在线/离线/新发现」;",
    "Telemetry 更新设备在线时间、流量统计等动态字段,并写入历史表;",
    "Risk Engine 对每台在线设备运行规则,生成风险等级,触发时写入告警表;",
    '命中"自动阻断"规则的设备进入 Blocker 队列执行阻断动作;',
    "所有动作经 Audit 模块统一落盘;",
    "Dashboard 通过 REST API 和 WebSocket 拉取/推送数据展示给用户。",
]
for s in flow:
    add_numbered(s)

add_h("2.4 部署架构", 2)
add_p("推荐部署方式为单机一体化部署:一台接入局域网的 Linux 主机(树莓派、迷你 PC、或网关本身)运行后端服务,后端同时作为静态资源服务器托管 React 构建产物。需要执行 ARP 扫描和阻断时,服务以 root 或具备 CAP_NET_RAW / CAP_NET_ADMIN 能力的用户运行。")


# ================== 三、技术栈 ==================
add_h("三、技术栈与依赖", 1)

add_h("3.1 后端核心依赖", 2)
add_table(
    ["依赖", "版本建议", "用途"],
    [
        ["Python",       "3.10+",    "运行时"],
        ["FastAPI",      "^0.110",   "Web 框架 / REST API"],
        ["Uvicorn",      "^0.27",    "ASGI 服务器"],
        ["SQLAlchemy",   "^2.0",     "ORM"],
        ["Pydantic",     "^2.6",     "数据模型与校验"],
        ["Scapy",        "^2.5",     "ARP 扫描、包构造、嗅探"],
        ["python-nmap",  "^0.7",     "端口探测(可选)"],
        ["APScheduler",  "^3.10",    "定时任务调度"],
        ["websockets",   "^12.0",    "实时推送"],
        ["manuf",        "^1.1",     "MAC OUI 厂商识别"],
        ["loguru",       "^0.7",     "日志"],
        ["pytest",       "^8.0",     "单元测试"],
    ],
)

add_h("3.2 前端核心依赖", 2)
add_table(
    ["依赖", "版本建议", "用途"],
    [
        ["React",       "^18.2",  "UI 框架"],
        ["Ant Design",  "^5.14",  "组件库"],
        ["Vite",        "^5.0",   "构建工具"],
        ["Axios",       "^1.6",   "HTTP 客户端"],
        ["ECharts",     "^5.5",   "图表"],
        ["Zustand",     "^4.5",   "状态管理(轻量)"],
        ["dayjs",       "^1.11",  "时间处理"],
    ],
)


# ================== 四、文件结构 ==================
add_h("四、文件结构总体框架", 1)
add_p("完整的项目目录结构如下:")
add_code("""device-access-control/
├── README.md
├── docker-compose.yml
├── .env.example
│
├── backend/                           # 后端 (Python / FastAPI)
│   ├── pyproject.toml
│   ├── requirements.txt
│   ├── README.md
│   ├── app/
│   │   ├── main.py                    # FastAPI 入口
│   │   ├── config.py                  # 配置加载
│   │   ├── deps.py
│   │   ├── api/                       # 路由层
│   │   │   ├── routes_devices.py
│   │   │   ├── routes_scan.py
│   │   │   ├── routes_risk.py
│   │   │   ├── routes_blocker.py
│   │   │   ├── routes_audit.py
│   │   │   └── routes_ws.py
│   │   ├── core/                      # 核心业务
│   │   │   ├── discovery.py
│   │   │   ├── fingerprint.py
│   │   │   ├── registry.py
│   │   │   ├── telemetry.py
│   │   │   ├── risk_engine.py
│   │   │   ├── blocker.py
│   │   │   ├── audit.py
│   │   │   └── scheduler.py
│   │   ├── models/                    # ORM 模型
│   │   │   ├── base.py
│   │   │   ├── device.py
│   │   │   ├── access_log.py
│   │   │   ├── alert.py
│   │   │   ├── rule.py
│   │   │   └── audit_log.py
│   │   ├── schemas/                   # Pydantic 数据模型
│   │   ├── utils/                     # 工具函数 (net/oui/logger/security)
│   │   └── ws/                        # WebSocket 连接管理
│   ├── data/                          # 运行时数据 (SQLite + OUI)
│   ├── scripts/                       # 运维脚本
│   │   ├── init_db.py
│   │   ├── seed_rules.py
│   │   └── export_audit.py
│   └── tests/                         # 单元测试
│
├── frontend/                          # 前端 (React + Ant Design)
│   ├── package.json
│   ├── vite.config.ts
│   ├── tsconfig.json
│   ├── index.html
│   └── src/
│       ├── main.tsx
│       ├── App.tsx
│       ├── router.tsx
│       ├── api/                       # API 封装
│       ├── pages/                     # Dashboard / Devices / Alerts / Rules / Audit / Settings
│       ├── components/
│       ├── hooks/
│       ├── store/                     # zustand
│       └── styles/
│
└── docs/                              # 设计文档
    ├── architecture.svg
    ├── data-model.md
    ├── api-spec.md
    └── deployment.md""")


# ================== 五、核心数据模型 ==================
add_h("五、核心数据模型", 1)

add_h("5.1 device 表(设备资产)", 2)
add_table(
    ["字段", "类型", "说明"],
    [
        ["id",         "INTEGER PK",  "自增主键"],
        ["mac",        "TEXT UNIQUE", "MAC 地址(主标识)"],
        ["ip",         "TEXT",        "当前 IP"],
        ["hostname",   "TEXT",        "主机名"],
        ["vendor",     "TEXT",        "OUI 厂商"],
        ["os_guess",   "TEXT",        "操作系统猜测"],
        ["category",   "TEXT",        "white / black / unknown"],
        ["risk_level", "INTEGER",     "0=低 1=中 2=高 3=严重"],
        ["status",     "TEXT",        "online / offline / blocked"],
        ["first_seen", "DATETIME",    "首次发现"],
        ["last_seen",  "DATETIME",    "最近在线"],
        ["note",       "TEXT",        "备注"],
    ],
)

add_h("5.2 其他核心表", 2)
add_p("access_log 表(接入历史):记录每次上线/下线/IP 变化事件,字段包括 id、device_id、event_type、ip、timestamp。")
add_p("alert 表(告警):记录风险引擎产生的告警,字段包括 id、device_id、rule_id、level、message、status (open/acknowledged/closed)、created_at。")
add_p("rule 表(规则):记录风险判定规则,字段包括 id、name、condition_json、action (alert/block)、level、enabled。")
add_p("audit_log 表(审计):记录所有人工/自动操作,字段包括 id、actor、action、target_device_id、reason、timestamp。")


# ================== 六、API 设计 ==================
add_h("六、API 设计概览", 1)
add_p("后端对外暴露 RESTful API,统一前缀 /api/v1,主要路由分组如下:")
api_groups = [
    ("设备相关",  "GET /devices(分页列表+筛选)、GET /devices/{id}、PATCH /devices/{id}、POST /devices/{id}/whitelist、POST /devices/{id}/blacklist。"),
    ("扫描相关",  "POST /scan/trigger(手动触发一次扫描)、GET /scan/status(当前扫描进度)。"),
    ("告警相关",  "GET /alerts、POST /alerts/{id}/ack。"),
    ("规则相关",  "GET /rules、POST /rules、PUT /rules/{id}、DELETE /rules/{id}。"),
    ("阻断相关",  "POST /blocker/{device_id}/block、POST /blocker/{device_id}/unblock、GET /blocker/active。"),
    ("审计相关",  "GET /audit、GET /audit/export。"),
    ("WebSocket", "WS /ws/realtime 推送设备上下线、新告警、阻断动作等实时事件。"),
]
for name, desc in api_groups:
    p = doc.add_paragraph()
    r1 = p.add_run(name + " —— ")
    set_cn_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_cn_font(r2)


# ================== 七、任务分解 WBS ==================
add_h("七、任务需求与工作分解 (WBS)", 1)
add_p("将整个项目划分为 5 个阶段、共约 28 个子任务,建议总工期 4~6 周(单人)或 2~3 周(2~3 人小组)。")

add_h("阶段一:环境与脚手架(预计 2 天)", 3)
add_table(
    ["编号", "任务", "产出"],
    [
        ["T1.1", "项目初始化", "仓库 + 目录结构 + .gitignore + README"],
        ["T1.2", "后端脚手架", "FastAPI 最小可运行 + 健康检查"],
        ["T1.3", "前端脚手架", "Vite + React + Ant Design 跑通空白布局"],
    ],
)

add_h("阶段二:核心功能开发(预计 10~15 天)", 3)
add_table(
    ["编号", "任务", "关键点"],
    [
        ["T2.1",  "数据库与 ORM",     "5 张表 + init_db.py"],
        ["T2.2",  "网卡 / 网段识别",   "utils/net.py 自动识别 LAN"],
        ["T2.3",  "ARP 扫描",          "Scapy srp() 实现"],
        ["T2.4",  "ICMP / Nmap 探测",  "补充活性检测"],
        ["T2.5",  "OUI 厂商识别",      "manuf 库集成"],
        ["T2.6",  "主机名采集",        "mDNS / NetBIOS / DHCP"],
        ["T2.7",  "设备注册与分类",     "白/黑/未知三态切换"],
        ["T2.8",  "状态采集与持久化",   "last_seen + access_log"],
        ["T2.9",  "风险规则引擎",      "JSON DSL + 5 条默认规则"],
        ["T2.10", "ARP 阻断实现",      "ARP 欺骗持续发送"],
        ["T2.11", "iptables 阻断(可选)", "网关部署模式"],
        ["T2.12", "审计日志",          "所有写操作统一封装"],
    ],
)

add_h("阶段三:接口与前端(预计 7~10 天)", 3)
add_table(
    ["编号", "任务", "关键点"],
    [
        ["T3.1", "设备相关 API",       "6 个路由,含分页/筛选/排序"],
        ["T3.2", "告警/规则/阻断/审计 API", "剩余 REST 接口"],
        ["T3.3", "WebSocket 推送",     "连接管理 + 广播"],
        ["T3.4", "前端布局与路由",      "侧边栏 + 顶栏 + 6 页面"],
        ["T3.5", "仪表盘页",            "统计 + 饼图 + 最新告警"],
        ["T3.6", "设备管理页",          "表格 + 详情抽屉"],
        ["T3.7", "告警 / 审计页",       "列表 + 筛选"],
        ["T3.8", "规则配置页",          "可视化编辑 condition_json"],
    ],
)

add_h("阶段四:测试与加固(预计 3~4 天)", 3)
add_table(
    ["编号", "任务", "关键点"],
    [
        ["T4.1", "单元测试", "discovery/fingerprint/risk_engine/blocker"],
        ["T4.2", "集成测试", "模拟设备上下线 + WebSocket"],
        ["T4.3", "安全加固", "鉴权/输入校验/二次确认/最小权限"],
    ],
)

add_h("阶段五:部署与交付(预计 2 天)", 3)
add_table(
    ["编号", "任务", "关键点"],
    [
        ["T5.1", "部署文档与脚本", "systemd / Dockerfile / docker-compose"],
        ["T5.2", "演示与验收",      "演示视频 + 实验报告"],
    ],
)


# ================== 八、风险点 ==================
add_h("八、风险点与注意事项", 1)
risks = [
    ("权限与合规", "ARP 扫描和阻断需要 root / CAP_NET_RAW / CAP_NET_ADMIN,部署时务必使用专用账号并最小化授权;在他人网络中使用本系统须取得授权,避免触及法律红线。"),
    ("ARP 阻断的副作用", "ARP 欺骗会持续向目标设备广播伪造响应,可能误伤同网段其他设备,且在严格的交换机(DAI 防护)下失效。建议在生产部署时优先使用网关 iptables 方案。"),
    ("MAC 伪造规避", "攻击者可能伪造白名单设备 MAC,系统应记录 MAC 突然出现在异常时段、与已知设备指纹不一致等可疑模式。"),
    ("扫描频率与噪声", "ARP 扫描过于频繁会产生网络噪声并被 IDS 告警,默认 30s 一次较为合理,可在配置中调整。"),
    ("时钟一致性", "所有时间戳使用 UTC 存储,展示时按浏览器时区转换。"),
]
for name, desc in risks:
    p = doc.add_paragraph()
    r1 = p.add_run(name + " —— ")
    set_cn_font(r1, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    r2 = p.add_run(desc)
    set_cn_font(r2)


# ================== 九、验收标准 ==================
add_h("九、验收标准", 1)
acceptance = [
    "能够在 1 分钟内发现新接入设备并写入数据库;",
    "白名单设备永不触发自动阻断;",
    "黑名单设备接入后 5 秒内被阻断并产生审计记录;",
    "界面上的设备列表、告警、审计可正常展示并支持筛选;",
    "所有写操作都能在审计日志中找到对应记录;",
    "单元测试覆盖率不低于 60%;",
    "具备完整的部署手册和演示材料。",
]
for s in acceptance:
    add_bullet(s)


# ================== 十、扩展方向 ==================
add_h("十、后续可扩展方向", 1)
add_p("完成基线功能后,可在以下方向继续扩展:基于 sklearn 的简单异常检测(替代纯规则);接入 ntop/Suricata 等深度流量分析;移动端 H5 告警推送;多网段 / 跨 VLAN 支持;与企业微信 / 钉钉 / 邮件告警渠道集成;基于 OpenWRT 的固件集成,真正在路由器上运行。")

# 结尾
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run("— 文档结束 —")
set_cn_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))

doc.save(OUT)
print(f"✅ Word 文档已生成: {OUT}")
